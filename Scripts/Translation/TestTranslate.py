import os
from pdf2docx import Converter
from deep_translator import GoogleTranslator
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox

def select_pdf_file():
    """Opens a file dialog to select a PDF file."""
    file_path = filedialog.askopenfilename(title="Select a PDF file", filetypes=[("PDF Files", "*.pdf")])
    return file_path

def convert_pdf_to_docx(pdf_path, docx_path):
    """Converts a PDF file to a DOCX file while keeping layout."""
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)  # Converts all pages
    cv.close()

def translate_docx(docx_path, target_language="en"):
    """Translates all text in a DOCX file while keeping layout."""
    doc = Document(docx_path)
    translator = GoogleTranslator(source="auto", target=target_language)
    
    for para in doc.paragraphs:
        if para.text.strip():  # Ensure it's not empty
            para.text = translator.translate(para.text)

    # Translate Table Contents
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = translator.translate(cell.text)

    doc.save(docx_path)

def convert_docx_to_pdf(docx_path, pdf_path):
    """Converts a DOCX file back to a PDF while keeping layout."""
    os.system(f'libreoffice --headless --convert-to pdf "{docx_path}" --outdir "{os.path.dirname(pdf_path)}"')

def recreate_pdf_with_translation(pdf_path):
    """Runs the full translation process while keeping the original layout."""
    docx_path = pdf_path.replace(".pdf", "_translated.docx")
    output_pdf = pdf_path.replace(".pdf", "_translated.pdf")

    convert_pdf_to_docx(pdf_path, docx_path)  # Convert PDF → DOCX
    translate_docx(docx_path)                 # Translate DOCX (Including Tables)
    convert_docx_to_pdf(docx_path, output_pdf)  # Convert back DOCX → PDF

    print(f"Translated PDF saved as: {output_pdf}")

def main():
    """Main function for user interaction."""
    root = tk.Tk()
    root.withdraw()
    pdf_path = select_pdf_file()
    if not pdf_path:
        messagebox.showerror("Error", "No file selected!")
        return

    recreate_pdf_with_translation(pdf_path)
    messagebox.showinfo("Success", f"Translated PDF saved at: {pdf_path.replace('.pdf', '_translated.pdf')}")

if __name__ == "__main__":
    main()
